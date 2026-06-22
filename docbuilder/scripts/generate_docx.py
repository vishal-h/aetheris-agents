import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from _format import format_cell

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
}


def _write_cell(doc_cell, cell_spec, col_type=None):
    para = doc_cell.paragraphs[0]
    run = para.add_run(format_cell(cell_spec["value"], col_type))
    run.bold = cell_spec["bold"]
    para.alignment = ALIGN_MAP.get(cell_spec["align"], WD_ALIGN_PARAGRAPH.LEFT)


def _add_section_heading(doc, text):
    """Add a sheet heading. A minimal base file may not define the 'Heading 1'
    style, so fall back to a bold paragraph rather than crashing."""
    try:
        doc.add_heading(text, level=1)
    except KeyError:
        para = doc.add_paragraph()
        para.add_run(text).bold = True


def _apply_table_style(table, style_name):
    """Apply the named table style. A minimal base file may not define it
    (e.g. 'Table Grid'); warn and leave the table unstyled rather than crash."""
    try:
        table.style = style_name
    except KeyError:
        print(
            json.dumps({
                "status": "warning",
                "warning": f"table style '{style_name}' not in document; using default",
            }),
            file=sys.stderr,
        )


def generate_docx(doc_spec, output_path, base_file=None):
    """Render the doc spec to docx.

    When ``base_file`` is given, the document is opened from it and content
    (sheet headings + tables) is appended after the existing body — the base
    file's header, footer, cover page, and styles are preserved. The top-level
    document title heading is skipped in base-file mode because the base file
    owns the cover/title. When ``base_file`` is None, a fresh document is
    created (m1 behaviour, unchanged).

    ``table_style`` is read from the doc spec (top-level, default "Table Grid")
    and applied per table. Note: ``header_row`` is xlsx-specific positioning and
    is intentionally not used here — the docx renderer builds tables from the
    logical ``rows`` array in order.
    """
    base_mode = base_file is not None
    doc = Document(base_file) if base_mode else Document()

    table_style = doc_spec.get("table_style", "Table Grid")

    if not base_mode:
        title = doc_spec.get("title")
        if title:
            doc.add_heading(title, level=0)

    for sheet in doc_spec["sheets"]:
        rows = sheet["rows"]
        columns = sheet["columns"]

        if not rows:
            continue

        _add_section_heading(doc, sheet["name"])

        table = doc.add_table(rows=len(rows), cols=len(columns))
        _apply_table_style(table, table_style)

        for row_idx, row_spec in enumerate(rows):
            tr = table.rows[row_idx]
            is_header = row_spec["type"] == "header"
            for col_idx, cell_spec in enumerate(row_spec["cells"]):
                # Header cells hold the column name; data/aggregate cells are
                # formatted by their column's type (currency/number).
                col_type = None if is_header else (
                    columns[col_idx]["type"] if col_idx < len(columns) else None
                )
                _write_cell(tr.cells[col_idx], cell_spec, col_type)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", default="output")
    parser.add_argument("--filename", default="document")
    parser.add_argument("--input", default=None, help="doc spec JSON file path (default: stdin)")
    parser.add_argument("--base-file", default=None,
                        help="docx base file to open and append into (branding preserved); "
                             "default: create a fresh document")
    args = parser.parse_args()

    try:
        src = open(args.input, encoding="utf-8") if args.input else sys.stdin
        doc_spec = json.load(src)
    except Exception as e:
        print(json.dumps({"status": "error", "error": str(e)}), file=sys.stderr)
        sys.exit(1)

    try:
        output_dir = Path(args.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{args.filename}.docx"
        generate_docx(doc_spec, output_path, base_file=args.base_file)
        print(str(output_path))
    except Exception as e:
        print(json.dumps({"status": "error", "error": str(e)}), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
