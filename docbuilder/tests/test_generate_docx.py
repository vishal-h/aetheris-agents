import json
import subprocess
import sys
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from generate_docx import generate_docx

USE_CASE_ROOT = Path(__file__).parent.parent
DEMO_DOCX = USE_CASE_ROOT / "data" / "templates" / "demo" / "proposal_v1.docx"


# --- fixture ---

@pytest.fixture
def simple_spec():
    return {
        "title": "Test Proposal",
        "template_id": "test/v1",
        "output_formats": ["docx"],
        "sheets": [
            {
                "name": "Line Items",
                "header_row": 2,
                "columns": [
                    {"name": "Code",  "type": "string",   "width": 12},
                    {"name": "Qty",   "type": "number",   "width": 8},
                    {"name": "Total", "type": "currency", "width": 10},
                ],
                "merge_ranges": [],
                "rows": [
                    {
                        "type": "header",
                        "cells": [
                            {"value": "Code",  "bold": True, "align": "left"},
                            {"value": "Qty",   "bold": True, "align": "right"},
                            {"value": "Total", "bold": True, "align": "right"},
                        ],
                    },
                    {
                        "type": "data",
                        "cells": [
                            {"value": "A-01",   "bold": False, "align": "left"},
                            {"value": "5",      "bold": False, "align": "right"},
                            {"value": "250.00", "bold": False, "align": "right"},
                        ],
                    },
                    {
                        "type": "data",
                        "cells": [
                            {"value": "B-02",   "bold": False, "align": "left"},
                            {"value": "3",      "bold": False, "align": "right"},
                            {"value": "150.00", "bold": False, "align": "right"},
                        ],
                    },
                    {
                        "type": "aggregate",
                        "cells": [
                            {"value": "",    "bold": True, "align": "left"},
                            {"value": 8,     "bold": True, "align": "right"},
                            {"value": 400,   "bold": True, "align": "right"},
                        ],
                    },
                ],
            },
            {
                "name": "Summary",
                "header_row": 1,
                "columns": [
                    {"name": "Metric", "type": "string", "width": 18},
                    {"name": "Value",  "type": "string", "width": 12},
                ],
                "merge_ranges": [],
                "rows": [
                    {
                        "type": "header",
                        "cells": [
                            {"value": "Metric", "bold": True, "align": "left"},
                            {"value": "Value",  "bold": True, "align": "right"},
                        ],
                    },
                    {
                        "type": "data",
                        "cells": [
                            {"value": "Total Items", "bold": True,  "align": "left"},
                            {"value": 2,             "bold": False, "align": "right"},
                        ],
                    },
                ],
            },
        ],
    }


# --- unit tests ---

@pytest.mark.integration
def test_file_created(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    assert out.exists()
    assert out.stat().st_size > 0


@pytest.mark.integration
def test_document_title_present(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    titles = [p for p in doc.paragraphs if p.style.name == "Title"]
    assert len(titles) == 1
    assert titles[0].text == "Test Proposal"


@pytest.mark.integration
def test_table_count_matches_sheet_count(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    assert len(doc.tables) == 2


@pytest.mark.integration
def test_header_row_cell_text(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    first_table = doc.tables[0]
    assert first_table.rows[0].cells[0].text == "Code"
    assert first_table.rows[0].cells[1].text == "Qty"
    assert first_table.rows[0].cells[2].text == "Total"


@pytest.mark.integration
def test_data_row_cell_text(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    table = doc.tables[0]
    assert table.rows[1].cells[0].text == "A-01"
    assert table.rows[2].cells[0].text == "B-02"


@pytest.mark.integration
def test_row_count(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    # 1 header + 2 data + 1 aggregate = 4
    assert len(doc.tables[0].rows) == 4


@pytest.mark.integration
def test_header_cell_bold(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    cell = doc.tables[0].rows[0].cells[0]
    run = cell.paragraphs[0].runs[0]
    assert run.bold is True


@pytest.mark.integration
def test_data_cell_not_bold(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    cell = doc.tables[0].rows[1].cells[0]
    run = cell.paragraphs[0].runs[0]
    assert run.bold is False


@pytest.mark.integration
def test_aggregate_cell_bold(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    # aggregate row is the last row (index 3)
    cell = doc.tables[0].rows[3].cells[1]
    run = cell.paragraphs[0].runs[0]
    assert run.bold is True


@pytest.mark.integration
def test_right_align_applied(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    # header row, Qty cell (index 1) — align: right
    para = doc.tables[0].rows[0].cells[1].paragraphs[0]
    assert para.alignment == WD_ALIGN_PARAGRAPH.RIGHT


@pytest.mark.integration
def test_left_align_applied(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    # header row, Code cell (index 0) — align: left
    para = doc.tables[0].rows[0].cells[0].paragraphs[0]
    assert para.alignment == WD_ALIGN_PARAGRAPH.LEFT


@pytest.mark.integration
def test_numeric_value_written_as_string(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    # aggregate row, Qty cell = 8 (int) → written as "8"
    cell = doc.tables[0].rows[3].cells[1]
    assert cell.text == "8"


@pytest.mark.integration
def test_summary_table_row_count(tmp_path, simple_spec):
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    # Summary: 1 header + 1 data = 2
    assert len(doc.tables[1].rows) == 2


# --- CLI integration ---

@pytest.mark.integration
def test_cli_produces_file(tmp_path):
    fetch = subprocess.run(
        [sys.executable, "scripts/fetch_data.py", "data/sample_data.csv"],
        capture_output=True, text=True, cwd=str(USE_CASE_ROOT)
    )
    assert fetch.returncode == 0
    compute = subprocess.run(
        [sys.executable, "scripts/compute_doc.py",
         "data/templates/demo/proposal_v1.json", "-"],
        input=fetch.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    assert compute.returncode == 0
    render = subprocess.run(
        [sys.executable, "scripts/generate_docx.py",
         "--output-dir", str(tmp_path), "--filename", "proposal"],
        input=compute.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    assert render.returncode == 0
    assert (tmp_path / "proposal.docx").exists()


@pytest.mark.integration
def test_cli_prints_output_path(tmp_path):
    fetch = subprocess.run(
        [sys.executable, "scripts/fetch_data.py", "data/sample_data.csv"],
        capture_output=True, text=True, cwd=str(USE_CASE_ROOT)
    )
    compute = subprocess.run(
        [sys.executable, "scripts/compute_doc.py",
         "data/templates/demo/proposal_v1.json", "-"],
        input=fetch.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    render = subprocess.run(
        [sys.executable, "scripts/generate_docx.py",
         "--output-dir", str(tmp_path), "--filename", "out"],
        input=compute.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    assert render.returncode == 0
    assert render.stdout.strip().endswith("out.docx")


@pytest.mark.integration
def test_cli_proposal_two_tables(tmp_path):
    fetch = subprocess.run(
        [sys.executable, "scripts/fetch_data.py", "data/sample_data.csv"],
        capture_output=True, text=True, cwd=str(USE_CASE_ROOT)
    )
    compute = subprocess.run(
        [sys.executable, "scripts/compute_doc.py",
         "data/templates/demo/proposal_v1.json", "-"],
        input=fetch.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    render = subprocess.run(
        [sys.executable, "scripts/generate_docx.py",
         "--output-dir", str(tmp_path), "--filename", "prop"],
        input=compute.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    assert render.returncode == 0
    doc = Document(str(tmp_path / "prop.docx"))
    assert len(doc.tables) == 2


@pytest.mark.integration
def test_cli_proposal_line_items_row_count(tmp_path):
    fetch = subprocess.run(
        [sys.executable, "scripts/fetch_data.py", "data/sample_data.csv"],
        capture_output=True, text=True, cwd=str(USE_CASE_ROOT)
    )
    compute = subprocess.run(
        [sys.executable, "scripts/compute_doc.py",
         "data/templates/demo/proposal_v1.json", "-"],
        input=fetch.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    render = subprocess.run(
        [sys.executable, "scripts/generate_docx.py",
         "--output-dir", str(tmp_path), "--filename", "prop"],
        input=compute.stdout, capture_output=True, text=True,
        cwd=str(USE_CASE_ROOT)
    )
    doc = Document(str(tmp_path / "prop.docx"))
    table = doc.tables[0]
    # 1 header + 10 data + 1 aggregate = 12 rows
    assert len(table.rows) == 12
    assert table.rows[0].cells[0].text == "Item Code"
    assert table.rows[1].cells[0].text == "SRV-001"
    assert table.rows[11].cells[1].text == "TOTAL"


# --- table_style (m2a t3) ---

@pytest.mark.integration
def test_table_style_defaults_to_table_grid(tmp_path, simple_spec):
    # No table_style in doc spec → default "Table Grid".
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    assert doc.tables[0].style.name == "Table Grid"


@pytest.mark.integration
def test_table_style_from_doc_spec(tmp_path, simple_spec):
    # Explicit table_style is read from the doc spec and applied.
    simple_spec["table_style"] = "Light List Accent 1"
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    doc = Document(str(out))
    assert doc.tables[0].style.name == "Light List Accent 1"


@pytest.mark.integration
def test_unknown_table_style_degrades(tmp_path, simple_spec, capsys):
    # An unknown style must not crash — warn on stderr, leave default formatting.
    simple_spec["table_style"] = "No Such Style 9000"
    out = tmp_path / "out.docx"
    generate_docx(simple_spec, out)
    assert out.exists()
    err = capsys.readouterr().err
    assert "No Such Style 9000" in err


# --- base file support (m2a t3) ---

@pytest.mark.integration
def test_base_file_absent_creates_fresh_document(tmp_path, simple_spec):
    out = tmp_path / "fresh.docx"
    generate_docx(simple_spec, out, base_file=None)
    doc = Document(str(out))
    # fresh document carries the doc title heading
    titles = [p for p in doc.paragraphs if p.style.name == "Title"]
    assert len(titles) == 1


@pytest.mark.integration
def test_base_file_header_preserved(tmp_path, simple_spec):
    out = tmp_path / "branded.docx"
    generate_docx(simple_spec, out, base_file=str(DEMO_DOCX))
    doc = Document(str(out))
    header_text = doc.sections[0].header.paragraphs[0].text
    assert "[ LOGO ]" in header_text
    assert "Company Name" in header_text


@pytest.mark.integration
def test_base_file_tables_appended(tmp_path, simple_spec):
    out = tmp_path / "branded.docx"
    generate_docx(simple_spec, out, base_file=str(DEMO_DOCX))
    doc = Document(str(out))
    # two sheets → two tables appended to the base body
    assert len(doc.tables) == 2
    assert doc.tables[0].rows[0].cells[0].text == "Code"


@pytest.mark.integration
def test_base_file_title_not_duplicated(tmp_path, simple_spec):
    # Base file already carries the cover title; base mode must not add another.
    out = tmp_path / "branded.docx"
    generate_docx(simple_spec, out, base_file=str(DEMO_DOCX))
    doc = Document(str(out))
    # the fresh-mode title heading uses style "Title"; in base mode none is added
    title_headings = [p for p in doc.paragraphs if p.style and p.style.name == "Title"]
    assert title_headings == []


@pytest.mark.integration
def test_base_file_missing_styles_degrade(tmp_path, simple_spec):
    # The demo base file lacks 'Heading 1' and 'Table Grid' styles — rendering
    # must still succeed (defensive fallback), producing the tables.
    out = tmp_path / "branded.docx"
    generate_docx(simple_spec, out, base_file=str(DEMO_DOCX))
    assert out.exists() and out.stat().st_size > 0
    doc = Document(str(out))
    assert len(doc.tables) == 2


@pytest.mark.integration
def test_cli_base_file_flag(tmp_path):
    fetch = subprocess.run(
        [sys.executable, "scripts/fetch_data.py", "data/sample_data.csv"],
        capture_output=True, text=True, cwd=str(USE_CASE_ROOT)
    )
    compute = subprocess.run(
        [sys.executable, "scripts/compute_doc.py",
         "data/templates/demo/proposal_v1.json", "-"],
        input=fetch.stdout, capture_output=True, text=True, cwd=str(USE_CASE_ROOT)
    )
    render = subprocess.run(
        [sys.executable, "scripts/generate_docx.py",
         "--base-file", "data/templates/demo/proposal_v1.docx",
         "--output-dir", str(tmp_path), "--filename", "branded"],
        input=compute.stdout, capture_output=True, text=True, cwd=str(USE_CASE_ROOT)
    )
    assert render.returncode == 0, render.stderr
    doc = Document(str(tmp_path / "branded.docx"))
    assert len(doc.tables) == 2
    header_text = doc.sections[0].header.paragraphs[0].text
    assert "[ LOGO ]" in header_text
